"""
Document processing workflow for handling file uploads.
"""

import asyncio
from typing import Dict, Any, Optional, List
from datetime import datetime
import io

from ..core.logger import get_logger
from ..core.database import get_database_session
from ..ai.gemini_service import get_gemini_service
from ..models import Document, DocumentCreate, FileMetadata, FileMetadataCreate
from ..utils.file_utils import is_supported_file_type, get_file_extension, format_file_size
from ..utils.text_utils import clean_text, chunk_text
from ..core.exceptions import TelegramAssistantException

logger = get_logger(__name__)


class DocumentProcessingError(TelegramAssistantException):
    """Exception for document processing errors."""
    pass


class DocumentProcessor:
    """Handles document processing and analysis."""

    def __init__(self):
        self.ai_service = None

    async def _get_ai_service(self):
        """Get AI service instance."""
        if self.ai_service is None:
            self.ai_service = await get_gemini_service()
        return self.ai_service

    async def process_file(
        self,
        file_data: bytes,
        filename: str,
        file_type: str,
        source: str = "upload"
    ) -> Dict[str, Any]:
        """Process uploaded file and extract content."""
        try:
            logger.info(f"Processing file: {filename}", extra={
                "filename": filename,
                "file_type": file_type,
                "file_size": len(file_data),
                "source": source
            })

            # Validate file type
            if not is_supported_file_type(filename):
                return {
                    "success": False,
                    "error": f"Unsupported file type: {get_file_extension(filename)}"
                }

            # Check file size
            max_size = 50 * 1024 * 1024  # 50MB limit
            if len(file_data) > max_size:
                return {
                    "success": False,
                    "error": f"File too large. Maximum size: {format_file_size(max_size)}"
                }

            # Create file metadata
            file_metadata = await self._create_file_metadata(filename, file_type, len(file_data), source)

            # Extract text content
            text_content = await self._extract_text_content(file_data, filename, file_type)

            if not text_content:
                return {
                    "success": False,
                    "error": "No text content could be extracted from the file"
                }

            # Process with AI
            ai_service = await self._get_ai_service()
            analysis = await ai_service.analyze_document(text_content, file_type)

            # Create document record
            document = await self._create_document_record(
                file_metadata=file_metadata,
                text_content=text_content,
                analysis=analysis
            )

            # Store in vector database (placeholder)
            # await self._store_in_vector_db(document, text_content)

            logger.info(f"Successfully processed file: {filename}", extra={
                "document_id": document.id,
                "text_length": len(text_content),
                "keywords": analysis.get("keywords", [])
            })

            return {
                "success": True,
                "document_id": document.id,
                "summary": analysis.get("summary", "Document processed successfully"),
                "keywords": analysis.get("keywords", []),
                "insights": analysis.get("insights", []),
                "text_length": len(text_content)
            }

        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}", extra={
                "filename": filename,
                "file_type": file_type,
                "error_type": type(e).__name__
            })
            return {
                "success": False,
                "error": f"Processing failed: {str(e)}"
            }

    async def _extract_text_content(self, file_data: bytes, filename: str, file_type: str) -> str:
        """Extract text content from file based on type."""
        try:
            file_extension = get_file_extension(filename).lower()

            if file_extension in ['txt']:
                return await self._extract_from_text(file_data)
            elif file_extension in ['pdf']:
                return await self._extract_from_pdf(file_data)
            elif file_extension in ['docx']:
                return await self._extract_from_docx(file_data)
            elif file_extension in ['xlsx', 'csv']:
                return await self._extract_from_spreadsheet(file_data, file_extension)
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                return await self._extract_from_image(file_data, filename)
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_extension}")
                return ""

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return ""

    async def _extract_from_text(self, file_data: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1']:
                try:
                    text = file_data.decode(encoding)
                    return clean_text(text)
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            text = file_data.decode('utf-8', errors='replace')
            return clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting text from plain text file: {e}")
            return ""

    async def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF file."""
        try:
            # Try to use PyPDF2 if available
            try:
                import PyPDF2
                import io

                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                return clean_text(text)

            except ImportError:
                logger.warning("PyPDF2 not available, using AI for PDF processing")
                # Fallback: Use AI to analyze the PDF as binary data
                ai_service = await self._get_ai_service()
                return await ai_service.extract_text_from_image(file_data, "document.pdf")

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""

    async def _extract_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            # Try to use python-docx if available
            try:
                from docx import Document
                import io

                doc = Document(io.BytesIO(file_data))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                return clean_text(text)

            except ImportError:
                logger.warning("python-docx not available, treating as binary")
                return "DOCX file processed (text extraction requires python-docx library)"

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""

    async def _extract_from_spreadsheet(self, file_data: bytes, file_extension: str) -> str:
        """Extract text from spreadsheet file."""
        try:
            if file_extension == 'csv':
                # Handle CSV
                import csv
                import io

                text = file_data.decode('utf-8', errors='replace')
                csv_reader = csv.reader(io.StringIO(text))
                extracted_text = ""
                for row in csv_reader:
                    extracted_text += " | ".join(row) + "\n"

                return clean_text(extracted_text)

            else:
                # Handle Excel files
                try:
                    import pandas as pd
                    import io

                    excel_file = pd.ExcelFile(io.BytesIO(file_data))
                    extracted_text = ""

                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)
                        extracted_text += f"\n\nSheet: {sheet_name}\n"
                        extracted_text += df.to_string(index=False)

                    return clean_text(extracted_text)

                except ImportError:
                    logger.warning("pandas not available for Excel processing")
                    return "Excel file processed (full text extraction requires pandas library)"

        except Exception as e:
            logger.error(f"Error extracting text from spreadsheet: {e}")
            return ""

    async def _extract_from_image(self, file_data: bytes, filename: str) -> str:
        """Extract text from image using AI OCR."""
        try:
            ai_service = await self._get_ai_service()
            return await ai_service.extract_text_from_image(file_data, filename)

        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""

    async def _create_file_metadata(self, filename: str, file_type: str, file_size: int, source: str) -> FileMetadata:
        """Create file metadata record."""
        try:
            async with get_database_session() as db_session:
                file_metadata = FileMetadata(
                    file_id=f"file_{datetime.now().timestamp()}",
                    original_filename=filename,
                    file_extension=get_file_extension(filename),
                    file_size_bytes=file_size,
                    source=source,
                    status="processing"
                )

                db_session.add(file_metadata)
                await db_session.commit()
                await db_session.refresh(file_metadata)

                return file_metadata

        except Exception as e:
            logger.error(f"Error creating file metadata: {e}")
            raise DocumentProcessingError(f"Failed to create file metadata: {e}")

    async def _create_document_record(
        self,
        file_metadata: FileMetadata,
        text_content: str,
        analysis: Dict[str, Any]
    ) -> Document:
        """Create document record in database."""
        try:
            async with get_database_session() as db_session:
                document = Document(
                    file_id=file_metadata.file_id,
                    original_filename=file_metadata.original_filename,
                    file_type=file_metadata.file_extension,
                    file_size_bytes=file_metadata.file_size_bytes,
                    mime_type=f"application/{file_metadata.file_extension}",
                    extracted_text=text_content[:50000],  # Limit database storage
                    overarching_theme=analysis.get("summary", ""),
                    keywords=analysis.get("keywords", []),
                    processing_status="completed"
                )

                db_session.add(document)
                await db_session.commit()
                await db_session.refresh(document)

                # Update file metadata status
                file_metadata.status = "completed"
                await db_session.commit()

                return document

        except Exception as e:
            logger.error(f"Error creating document record: {e}")
            raise DocumentProcessingError(f"Failed to create document record: {e}")

    async def search_documents(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents in the database."""
        try:
            async with get_database_session() as db_session:
                from sqlalchemy import select, or_
                from ..models import Document

                # Simple text search (can be enhanced with vector search later)
                search_query = f"%{query.lower()}%"

                stmt = select(Document).where(
                    or_(
                        Document.extracted_text.ilike(search_query),
                        Document.overarching_theme.ilike(search_query),
                        Document.original_filename.ilike(search_query)
                    )
                ).limit(limit)

                result = await db_session.execute(stmt)
                documents = result.scalars().all()

                search_results = []
                for doc in documents:
                    # Find relevant snippet
                    text = doc.extracted_text or ""
                    query_pos = text.lower().find(query.lower())

                    if query_pos >= 0:
                        start = max(0, query_pos - 100)
                        end = min(len(text), query_pos + 200)
                        snippet = text[start:end]
                    else:
                        snippet = text[:200] if text else doc.overarching_theme

                    search_results.append({
                        "document_id": doc.id,
                        "filename": doc.original_filename,
                        "snippet": snippet,
                        "score": 0.8,  # Placeholder score
                        "keywords": doc.keywords or []
                    })

                return search_results

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []

    async def get_document_stats(self) -> Dict[str, Any]:
        """Get document processing statistics."""
        try:
            async with get_database_session() as db_session:
                from sqlalchemy import select, func
                from ..models import Document, FileMetadata

                # Count documents
                doc_count_result = await db_session.execute(select(func.count()).select_from(Document))
                total_documents = doc_count_result.scalar() or 0

                # Count by status
                processing_count_result = await db_session.execute(
                    select(func.count()).select_from(Document).where(
                        Document.processing_status == "processing"
                    )
                )
                processing_count = processing_count_result.scalar() or 0

                # Total file size
                size_result = await db_session.execute(
                    select(func.sum(Document.file_size_bytes)).select_from(Document)
                )
                total_size = size_result.scalar() or 0

                return {
                    "total_documents": total_documents,
                    "processing_count": processing_count,
                    "completed_count": total_documents - processing_count,
                    "total_size_bytes": total_size,
                    "total_size_formatted": format_file_size(total_size)
                }

        except Exception as e:
            logger.error(f"Error getting document stats: {e}")
            return {
                "total_documents": 0,
                "processing_count": 0,
                "completed_count": 0,
                "total_size_bytes": 0,
                "total_size_formatted": "0 B"
            }


# Global instance
document_processor = DocumentProcessor()